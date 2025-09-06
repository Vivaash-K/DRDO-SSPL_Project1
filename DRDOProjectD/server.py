from flask import Flask, jsonify, request, send_from_directory, send_file
from record_audio import record_audio, stop_recording
from transcribe_audio import transcribe_audio
from summarize_text import summarize_text
import os
from datetime import datetime
import threading
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil

app = Flask(__name__, static_folder='static')

# Global variable to track recording state
is_recording = False
recording_thread = None

@app.route('/')
def home():
    return send_from_directory('static', 'index.html')

@app.route('/start_recording', methods=['POST'])
def start_recording():
    global is_recording, recording_thread
    if not is_recording:
        is_recording = True
        try:
            # Start recording in a separate thread
            recording_thread = threading.Thread(target=record_audio)
            recording_thread.start()
            return jsonify({"status": "success", "message": "Recording started"})
        except Exception as e:
            is_recording = False
            return jsonify({"status": "error", "message": str(e)})
    else:
        is_recording = False
        try:
            # Stop the recording
            stop_recording()
            if recording_thread and recording_thread.is_alive():
                recording_thread.join(timeout=2)  # Wait for the recording to finish
            return jsonify({"status": "success", "message": "Recording stopped"})
        except Exception as e:
            return jsonify({"status": "error", "message": str(e)})

@app.route('/transcribe', methods=['POST'])
def transcribe():
    try:
        # Get the most recent recording
        recordings_dir = "recordings"
        if not os.path.exists(recordings_dir):
            return jsonify({"status": "error", "message": "No recordings found"})
        
        wav_files = [f for f in os.listdir(recordings_dir) if f.endswith('.wav')]
        if not wav_files:
            return jsonify({"status": "error", "message": "No WAV files found"})
        
        latest_recording = max(
            [os.path.join(recordings_dir, f) for f in wav_files],
            key=os.path.getmtime
        )
        
        # Transcribe the audio
        transcription_file = transcribe_audio(latest_recording)
        
        # Read the transcription
        with open(transcription_file, 'r', encoding='utf-8') as f:
            transcription = f.read()
        
        return jsonify({
            "status": "success",
            "transcription": transcription,
            "file": transcription_file
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})

@app.route('/summarize', methods=['POST'])
def summarize():
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({"status": "error", "message": "No text provided"})
        
        # Get the text to summarize
        text = data['text']
        
        # Generate summary
        summary_file = summarize_text(text)
        
        # Read the summary
        with open(summary_file, 'r', encoding='utf-8') as f:
            summary = f.read()
        
        return jsonify({
            "status": "success",
            "summary": summary,
            "file": summary_file
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})

@app.route('/download', methods=['POST'])
def download():
    try:
        data = request.get_json()
        if not data or 'transcription' not in data or 'summary' not in data:
            return jsonify({"status": "error", "message": "Missing transcription or summary"})
        
        # Create a new Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Meeting Summary', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        timestamp = doc.add_paragraph(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add transcription section
        doc.add_heading('Transcription', level=1)
        transcription_paragraph = doc.add_paragraph(data['transcription'])
        
        # Add summary section
        doc.add_heading('Summary', level=1)
        summary_paragraph = doc.add_paragraph(data['summary'])
        
        # Create documents directory if it doesn't exist
        if not os.path.exists('documents'):
            os.makedirs('documents')
        
        # Save the document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_path = f"documents/meeting_summary_{timestamp}.docx"
        doc.save(doc_path)
        
        # Send the file
        return send_file(
            doc_path,
            as_attachment=True,
            download_name=f"meeting_summary_{timestamp}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})

@app.route('/wipe_data', methods=['POST'])
def wipe_data():
    try:
        # List of directories to wipe
        directories = ['recordings', 'transcriptions', 'summaries']
        
        for directory in directories:
            if os.path.exists(directory):
                # Remove all files in the directory
                for filename in os.listdir(directory):
                    file_path = os.path.join(directory, filename)
                    try:
                        if os.path.isfile(file_path):
                            os.unlink(file_path)
                    except Exception as e:
                        return jsonify({"status": "error", "message": f"Error deleting {file_path}: {str(e)}"})
        
        return jsonify({"status": "success", "message": "All data has been wiped successfully"})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})

if __name__ == '__main__':
    # Create necessary directories if they don't exist
    os.makedirs('recordings', exist_ok=True)
    os.makedirs('transcriptions', exist_ok=True)
    os.makedirs('summaries', exist_ok=True)
    os.makedirs('documents', exist_ok=True)
    app.run(debug=True, port=5000) 